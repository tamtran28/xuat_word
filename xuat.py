import streamlit as st
import openpyxl
from docx import Document
from datetime import datetime
import tempfile
import os

st.set_page_config(page_title="Xuất bảng từ sheets chỉ định", page_icon="📄")
st.title("📄 Xuất dữ liệu từ các sheet được nhập trong Excel sang Word")

uploaded_file = st.file_uploader("🔽 Chọn file Excel (.xlsx hoặc .xlsm)", type=["xlsx", "xlsm"])
range_address = st.text_input("📌 Vùng bảng (VD: A1:G20)", value="A1:M103")

sheet_input = st.text_input("📑 Nhập tên sheets (cách nhau bởi dấu phẩy)", value="TK_KPCS_BANG_01, TK_KPCS_BANG_02")

if st.button("📤 Xuất sang Word"):
    if not uploaded_file:
        st.error("⚠️ Vui lòng tải lên file Excel.")
    elif sheet_input.strip() == "":
        st.error("⚠️ Bạn chưa nhập tên sheet.")
    else:
        try:
            # Xử lý danh sách sheet từ input
            selected_sheets = [s.strip() for s in sheet_input.split(",")]

            # Lưu file Excel tạm
            temp_excel = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsm")
            temp_excel.write(uploaded_file.read())
            temp_excel.close()

            # Mở Excel (có keep_vba)
            wb = openpyxl.load_workbook(temp_excel.name, data_only=True, keep_vba=True)

            # Tạo tài liệu Word
            doc = Document()
            doc.add_heading("📄 Dữ liệu từ các sheet được chọn", level=1)

            for sheet_name in selected_sheets:
                if sheet_name not in wb.sheetnames:
                    doc.add_paragraph(f"⚠️ Sheet không tồn tại: {sheet_name}")
                    continue

                ws = wb[sheet_name]
                doc.add_heading(f"📑 Sheet: {sheet_name}", level=2)

                try:
                    start_cell, end_cell = range_address.split(":")
                    data = ws[start_cell:end_cell]
                except Exception as e:
                    doc.add_paragraph(f"❌ Không thể đọc vùng {range_address} trên {sheet_name}: {e}")
                    continue

                # Tạo bảng Word
                table = doc.add_table(rows=len(data), cols=len(data[0]))
                table.style = "Table Grid"

                for i, row in enumerate(data):
                    for j, cell in enumerate(row):
                        value = cell.value if cell.value is not None else ""
                        table.cell(i, j).text = str(value)

                doc.add_paragraph()

            # Lưu Word tạm
            temp_word = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(temp_word.name)

            # Cho phép tải về
            with open(temp_word.name, "rb") as f:
                st.download_button(
                    label="📥 Tải file Word",
                    data=f.read(),
                    file_name=f"Export_{datetime.now():%Y%m%d_%H%M%S}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"❌ Lỗi: {e}")
